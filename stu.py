import streamlit as st
import json
import base64
from streamlit_option_menu import option_menu
import cohere
from io import BytesIO
from docx import Document
from api_call import get_api_key
from json_upload import upload_file

if not get_api_key():
    st.stop()
if not upload_file():
    st.stop()

cohere_api_key = st.session_state.api_key
co = cohere.Client(cohere_api_key)

def create_word_document(notes):
    doc = Document()
    doc.add_heading('Notes', 0)
    doc.add_paragraph(notes)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def decode_base64_data(encoded_data):
    try:
        return base64.b64decode(encoded_data)
    except (base64.binascii.Error, TypeError) as e:
        st.error("Failed to decode data, Upload the correct file.")
        st.stop()
        return None

def display_video(video_base64):
    video_bytes = decode_base64_data(video_base64)
    if video_bytes is not None:
        st.video(video_bytes)
    else:
        st.error("Error Video is not there, Upload the correct file.")
        st.stop()


def decode_and_download_files(uploaded_files):
    for file in uploaded_files:
        file_name = file['name']
        file_content = decode_base64_data(file['content'])

        st.download_button(
            label=f"Download {file_name}",
            data=file_content,
            file_name=file_name,
            mime='application/octet-stream'
        )

def parse_mcqs(mcq_list):
    parsed_mcqs = {}
    for item in mcq_list:
        try:
            difficulty, mcq_string = item.split(': ', 1)
            mcq_string = mcq_string.replace('\n', '').replace('\r', '').replace('    ', '').replace('easy', 'Easy:')\
                                    .replace('medium', 'Medium').replace('hard', 'Hard').replace('Id', 'id')\
                                    .replace('Question', 'question').replace('Options', 'options')\
                                    .replace('A)', 'a)').replace('B)', 'b)').replace('C)', 'c)')\
                                    .replace('D)', 'd)').replace('Answer', 'answer')
            parsed_mcqs[difficulty] = json.loads(mcq_string)
        except json.JSONDecodeError as e:
            st.error(f"Error parsing MCQs, Upload the correct file.")
            st.stop()
        except ValueError as e:
            st.error(f"Error splitting MCQs data, Upload the correct file.")
            st.stop()
    return parsed_mcqs

def parse_desc(desc_data):
    parsed_desc = {}
    try:
        desc_parts = desc_data.split('\n')
        for part in desc_parts:
            if ': ' in part:
                difficulty, questions = part.split(': ', 1)
                parsed_desc[difficulty] = json.loads(questions.replace('\n', '').replace('\r', ''))
    except json.JSONDecodeError as e:
        st.error(f"Error parsing Descriptive Questions, Upload the correct file.")
        st.stop()
    return parsed_desc

def display_mcqs(mcq_data, selected_difficulty):
    questions = mcq_data.get(selected_difficulty, [])
    answers = {}
    for q in questions:
        st.write(f"**{q['id']}. {q['question']}**")
        answers[q['id']] = st.radio("Options:", options=q['options'], key=q['id'])
    return answers

def display_desc(desc_data, selected_difficulty):
    questions = desc_data.get(selected_difficulty, [])
    responses = {}
    for q in questions:
        responses[q] = st.text_input(q, key=q)
    return responses

def evaluate_mcqs(mcq_answers, mcq_data, selected_difficulty, summary, notes):
    score = 0
    total_questions = 0
    questions = mcq_data.get(selected_difficulty, [])
    
    for q in questions:
        total_questions += 1
        correct_answer = q.get('answer')
        user_answer = mcq_answers.get(q['id'])
        if user_answer[0] == correct_answer[0]:
            st.success(f"{q['id']}. {q['question']}\n\nCorrect Answer: {correct_answer}")
            score += 1
        else:
            with st.spinner(f'Evaluating MCQ Question {q["id"]}'):
                prompt = f"""
                Question: {q['question']}
                Correct Answer: {correct_answer}
                
                Provide an explanation of why the correct answer is the best choice in very few lines (within 50 tokens).
                Summary: {summary}
                Notes: {notes}
                """
                response = co.generate(
                    model='command-light-nightly',
                    prompt=prompt,
                    max_tokens=100,
                    temperature=1.0,
                    k=0,
                    stop_sequences=["--"]
                )
                explanation = response.generations[0].text.strip()
                st.error(f"{q['id']}. {q['question']}\n\nUser Answer: {user_answer[0]} \n\nCorrect Answer: {correct_answer}\n\nExplanation: {explanation}")
                
    return score, total_questions

def evaluate_desc(desc_responses, summary, notes):
    evaluation_results = {}
    for question, user_answer in desc_responses.items():
        if user_answer:
            prompt = f"Evaluate the following answer based on the summary and notes:\n\nQuestion: {question}\nAnswer: {user_answer}\n\nSummary: {summary}\nNotes: {notes}\n\nProvide a score out of 5 and feedback. Give 0 score answer irrelevant to the summary and just say '0/5 (end of line) Answer is Irrelevant!'. If it is relevant, in range from 1 to 5, score how relevant it is. Mention the mistakes in the user's answer if any."
            response = co.generate(
                model='command-light-nightly',
                prompt=prompt,
            )
            evaluation_results[question] = response.generations[0].text.strip()
        else:
            evaluation_results[question] = "0/5\n\nAnswer the Question!"
    return evaluation_results

def chatbot_response(question, summary, transcript):
    prompt = f"Respond to the user's prompt briefly. Only If the user's prompt is related to given summary or transcript explain from that else don't mention anything about it. Respond only to the prompt. Don't put anything like 'Sure I will do that' or extra\n\nSummary: {summary}\n\nTranscript:{transcript}\n\nUser's prompt: {question}"
    response = co.generate(
        model='command-light-nightly',
        prompt=prompt,
        temperature=0.75
    )
    return response.generations[0].text.strip()

def chat_bot_ui(summary, transcript):
    messages = st.container()
    if prompt := st.chat_input("Say something"):
        messages.chat_message("user").write(prompt)
        messages.chat_message("assistant").write(chatbot_response(prompt, summary, transcript))
        
def main():
    data = st.session_state.data_json
    title = data['title']
    transcript = data['transcript'] if 'transcript' in data and data['transcript'] else ""
    summary = data['summary']
    notes = data['notes']
    quiz_data = data['quiz']
    video_base64 = data['video']
    uploaded_files = data['uploaded_files']
    
    st.title(title)

    display_video(video_base64)
    selected_menu = option_menu(
        menu_title="Menu",
        options=["Description", "Quiz", "Chatbot"],
        icons=["book", "pencil", "chat"],
        orientation="horizontal"
    )

    if selected_menu == "Description":
        with st.expander("# Summary:", expanded=True):
            st.write(summary)
        
        with st.expander("# Notes:"):
            st.write(notes)
            word_doc = create_word_document(notes)
            st.download_button(
                label="Download Notes as Word Document",
                data=word_doc,
                file_name="notes.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with st.expander("# Additional Notes:"):
            decode_and_download_files(uploaded_files)

    elif selected_menu == "Quiz":
        selected_mode = st.selectbox("Mode:", ("Easy", "Medium", "Hard"))
        
        st.write('#### MCQ Quiz')
        mcq_data = parse_mcqs(quiz_data['MCQ'])
        mcq_answers = display_mcqs(mcq_data, selected_mode)
        
        st.write('#### Descriptive Questions')
        desc_data = parse_desc(quiz_data['Desc'][0])
        desc_responses = display_desc(desc_data, selected_mode)
        
        if st.button('Submit All Answers'):
            st.title('Evaluation:')
            st.write("### MCQ:")
            mcq_score, mcq_total = evaluate_mcqs(mcq_answers, mcq_data, selected_mode, summary, notes)
            st.write(f"#### **MCQ Score:** {mcq_score}/{mcq_total}")
            
            with st.spinner('Evaluating Descriptive Questions'):
                desc_results = evaluate_desc(desc_responses, summary, notes)
            st.write("### Descriptive:")
            for question, feedback in desc_results.items():
                st.write(f"##### {question}")
                st.info(f"Feedback: {feedback}")

    elif selected_menu == "Chatbot":
        chat_bot_ui(summary, transcript)

if __name__ == '__main__':
    main()
